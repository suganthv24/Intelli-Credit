import io
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches

def generate_pdf(data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", "B", 16)
    
    # Title
    pdf.cell(0, 10, "Intelli-Credit Appraisal Memo", ln=True, align="C")
    pdf.ln(10)
    
    # Recommendation
    rec = data.get("recommendation", {})
    pdf.set_font("helvetica", "B", 12)
    pdf.cell(0, 10, "Credit Decision:", ln=True)
    pdf.set_font("helvetica", "", 12)
    pdf.cell(0, 8, f"Status: {rec.get('loan_decision', 'N/A')}", ln=True)
    pdf.cell(0, 8, f"Recommended Limit: Rs. {rec.get('recommended_limit', 0):,}", ln=True)
    pdf.cell(0, 8, f"Interest Rate: {rec.get('interest_rate', 0)}%", ln=True)
    pdf.ln(5)
    
    # Risk Profile
    pdf.set_font("helvetica", "B", 12)
    pdf.cell(0, 10, "Risk Assessment:", ln=True)
    pdf.set_font("helvetica", "", 12)
    pdf.cell(0, 8, f"Probability of Default: {rec.get('risk_probability', 0):.2%}", ln=True)
    pdf.cell(0, 8, f"Fraud Status: {data.get('fraud_status', 'CLEAN')}", ln=True)
    pdf.ln(5)
    
    # Due Diligence Notes
    diligence = data.get("diligence", {})
    if diligence:
        pdf.set_font("helvetica", "B", 12)
        pdf.cell(0, 10, "Due Diligence Notes:", ln=True)
        pdf.set_font("helvetica", "", 12)
        pdf.cell(0, 8, f"Factory Utilization: {diligence.get('utilization', 'N/A')}%", ln=True)
        pdf.cell(0, 8, f"Inventory Status: {diligence.get('inventory', 'N/A')}", ln=True)
        pdf.multi_cell(0, 8, f"Observations: {diligence.get('notes', 'None')}")
        pdf.ln(5)

    pdf_out = pdf.output(dest='S')
    if isinstance(pdf_out, str):
        pdf_bytes = pdf_out.encode('latin-1')
    else:
        pdf_bytes = bytearray(pdf_out)
        
    buffer = io.BytesIO(pdf_bytes)
    return buffer

def generate_docx(data):
    doc = Document()
    doc.add_heading("Intelli-Credit Appraisal Memo", 0)
    
    # Recommendation
    doc.add_heading("Credit Decision", level=1)
    rec = data.get("recommendation", {})
    p = doc.add_paragraph()
    p.add_run("Status: ").bold = True
    p.add_run(f"{rec.get('loan_decision', 'N/A')}\n")
    p.add_run("Recommended Limit: ").bold = True
    p.add_run(f"Rs. {rec.get('recommended_limit', 0):,}\n")
    p.add_run("Interest Rate: ").bold = True
    p.add_run(f"{rec.get('interest_rate', 0)}%")
    
    # Risk Profile
    doc.add_heading("Risk Assessment", level=1)
    p2 = doc.add_paragraph()
    p2.add_run("Probability of Default: ").bold = True
    p2.add_run(f"{rec.get('risk_probability', 0):.2%}\n")
    p2.add_run("Fraud Status: ").bold = True
    p2.add_run(f"{data.get('fraud_status', 'CLEAN')}")
    
    # Due Diligence Notes
    diligence = data.get("diligence", {})
    if diligence:
        doc.add_heading("Due Diligence Notes", level=1)
        p3 = doc.add_paragraph()
        p3.add_run("Factory Utilization: ").bold = True
        p3.add_run(f"{diligence.get('utilization', 'N/A')}%\n")
        p3.add_run("Inventory Status: ").bold = True
        p3.add_run(f"{diligence.get('inventory', 'N/A')}\n")
        p3.add_run("Observations: ").bold = True
        p3.add_run(f"{diligence.get('notes', 'None')}")
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
