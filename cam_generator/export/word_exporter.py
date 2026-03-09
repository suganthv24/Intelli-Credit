import os
from docx import Document

def generate_word_cam(sections, output_dir="."):
    """
    Given a dictionary of sections {title: content}, generates a formatted .docx report.
    """
    doc = Document()
    
    doc.add_heading("CREDIT APPRAISAL MEMO", level=0)
    
    for title, content in sections.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
        
    output_path = os.path.join(output_dir, "Credit_Appraisal_Memo.docx")
    doc.save(output_path)
    return output_path
